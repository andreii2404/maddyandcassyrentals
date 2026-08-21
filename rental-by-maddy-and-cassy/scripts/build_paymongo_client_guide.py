from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "PayMongo_Client_Configuration_Guide.docx"
LOGO = ROOT / "public" / "images" / "maddy-cassy-rentals-logo.png"

INK = "172033"
ROSE = "A85766"
ROSE_DARK = "8F4554"
BLUSH = "F9E8E6"
BLUSH_LIGHT = "FFF7F5"
GREEN = "176B45"
GREEN_LIGHT = "E8F5EE"
GOLD = "B97914"
GRAY = "5E6675"
LINE = "E7D9D6"
WHITE = "FFFFFF"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = "left" if edge == "start" else "right" if edge == "end" else edge
        el = borders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            borders.append(el)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                el.set(qn(f"w:{key}"), str(edge_data[key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    elem = p_pr.find(qn("w:keepNext"))
    if value and elem is None:
        elem = OxmlElement("w:keepNext")
        p_pr.append(elem)


def keep_lines(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    elem = p_pr.find(qn("w:keepLines"))
    if value and elem is None:
        elem = OxmlElement("w:keepLines")
        p_pr.append(elem)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_hyperlink(paragraph, text, url, color=ROSE_DARK, bold=False):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, fld_end])


def set_run(run, *, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def add_bullet(doc, text, *, level=0, bold_lead=None):
    p = doc.add_paragraph(style="Compact Bullet" if level == 0 else "Compact Bullet 2")
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=INK)
        set_run(p.add_run(text[len(bold_lead):]), color=INK)
    else:
        set_run(p.add_run(text), color=INK)
    keep_lines(p)
    return p


def add_numbered(doc, title, text):
    p = doc.add_paragraph(style="Compact Number")
    set_run(p.add_run(title), bold=True, color=INK)
    set_run(p.add_run(f" — {text}"), color=INK)
    keep_lines(p)
    return p


def add_callout(doc, title, body, fill=BLUSH_LIGHT, accent=ROSE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    border = {"val": "single", "sz": 8, "space": 0, "color": accent}
    set_cell_border(cell, top=border, bottom=border, start=border, end=border)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title), bold=True, color=accent, size=10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run(p2.add_run(body), color=INK, size=9.5)
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Title", 24, INK, 0, 5),
        ("Subtitle", 11, GRAY, 0, 10),
        ("Heading 1", 16, ROSE_DARK, 14, 7),
        ("Heading 2", 12.5, ROSE_DARK, 10, 5),
        ("Heading 3", 10.5, INK, 8, 3),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = styles.add_style("Compact Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = styles["Normal"]
    bullet.paragraph_format.left_indent = Inches(0.28)
    bullet.paragraph_format.first_line_indent = Inches(-0.18)
    bullet.paragraph_format.space_after = Pt(3)
    bullet.paragraph_format.line_spacing = 1.1
    bullet._element.get_or_add_pPr().append(OxmlElement("w:keepLines"))

    bullet2 = styles.add_style("Compact Bullet 2", WD_STYLE_TYPE.PARAGRAPH)
    bullet2.base_style = styles["Normal"]
    bullet2.paragraph_format.left_indent = Inches(0.52)
    bullet2.paragraph_format.first_line_indent = Inches(-0.18)
    bullet2.paragraph_format.space_after = Pt(2)

    number = styles.add_style("Compact Number", WD_STYLE_TYPE.PARAGRAPH)
    number.base_style = styles["Normal"]
    number.paragraph_format.left_indent = Inches(0.28)
    number.paragraph_format.first_line_indent = Inches(-0.2)
    number.paragraph_format.space_after = Pt(4)

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.base_style = styles["Normal"]
    code.font.name = "Consolas"
    code.font.size = Pt(8.7)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.14)
    code.paragraph_format.right_indent = Inches(0.14)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Rental by Maddy & Cassy  •  PayMongo Setup Guide  •  "), size=8, color=GRAY)
    add_field(p, "PAGE")
    for r in p.runs:
        set_run(r, size=8, color=GRAY)


def add_doc_header(section):
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(1.7)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    for cell in (left, right):
        set_cell_margins(cell, top=0, start=0, bottom=35, end=0)
    p = left.paragraphs[0]
    set_run(p.add_run("RENTAL BY MADDY & CASSY"), size=8.5, bold=True, color=ROSE_DARK)
    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p2.add_run("CLIENT SETUP GUIDE"), size=8, bold=True, color=GRAY)
    bottom = {"val": "single", "sz": 8, "space": 0, "color": LINE}
    set_cell_border(left, bottom=bottom)
    set_cell_border(right, bottom=bottom)


def add_env_table(doc):
    rows = [
        ("NEXT_PUBLIC_PAYMONGO_PUBLIC_KEY", "pk_test_…", "Safe for the browser; use the test public key."),
        ("PAYMONGO_SECRET_KEY", "sk_test_…", "Private server key. Never send in chat or place in frontend code."),
        ("PAYMONGO_WEBHOOK_SECRET", "whsk_…", "Copy from the test webhook you created."),
        ("PAYMONGO_PAYMENT_METHODS", "gcash,card,qrph", "Only include methods enabled for the account."),
        ("NEXT_PUBLIC_APP_URL", "https://maddyandcassyrentals-nine.vercel.app", "Use the final production/custom domain if it changes."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.25, 1.75, 2.5]
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    headers = ["Vercel variable", "Test value format", "What it controls"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, ROSE_DARK)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        set_run(p.add_run(text), bold=True, color=WHITE, size=9)
    set_repeat_table_header(table.rows[0])
    for row_no, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for col, text in enumerate(values):
            cell = row.cells[col]
            shade(cell, WHITE if row_no % 2 == 0 else BLUSH_LIGHT)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            font = "Consolas" if col in (0, 1) else "Calibri"
            set_run(p.add_run(text), size=8.5 if col in (0, 1) else 9, color=INK, font=font)
    border = {"val": "single", "sz": 4, "space": 0, "color": LINE}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)
    return table


def add_video(doc, number, title, url, why, caution=None):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.55)
    table.columns[1].width = Inches(5.95)
    badge = table.cell(0, 0)
    content = table.cell(0, 1)
    set_cell_width(badge, 0.55)
    set_cell_width(content, 5.95)
    shade(badge, ROSE_DARK)
    shade(content, BLUSH_LIGHT)
    for cell in (badge, content):
        set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
        border = {"val": "single", "sz": 5, "space": 0, "color": LINE}
        set_cell_border(cell, top=border, bottom=border, start=border, end=border)
    p = badge.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(str(number)), bold=True, color=WHITE, size=13)
    p = content.paragraphs[0]
    add_hyperlink(p, title, url, bold=True)
    p.paragraph_format.space_after = Pt(3)
    p2 = content.add_paragraph()
    p2.paragraph_format.space_after = Pt(0 if not caution else 2)
    set_run(p2.add_run(why), size=9.2, color=INK)
    if caution:
        p3 = content.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        set_run(p3.add_run(caution), size=8.5, color=GOLD, italic=True)
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.28)
    setup_styles(doc)
    add_doc_header(section)
    add_page_number_footer(section)

    title_table = doc.add_table(rows=1, cols=2)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_table.autofit = False
    title_table.columns[0].width = Inches(5.1)
    title_table.columns[1].width = Inches(1.4)
    left, right = title_table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=80, start=0, bottom=80, end=0)
    p = left.paragraphs[0]
    p.style = doc.styles["Title"]
    set_run(p.add_run("PayMongo Client Configuration Guide"), size=24, bold=True, color=INK)
    p2 = left.add_paragraph(style="Subtitle")
    set_run(p2.add_run("Dashboard setup, test payment, webhook, and safe live launch"), size=11, color=GRAY)
    p3 = left.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    set_run(p3.add_run("Prepared for: "), size=9, bold=True, color=ROSE_DARK)
    set_run(p3.add_run("Rental by Maddy & Cassy"), size=9, color=INK)
    if LOGO.exists():
        pr = right.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pr.add_run().add_picture(str(LOGO), width=Inches(1.05))

    add_callout(
        doc,
        "The short version",
        "Your client completes PayMongo verification, enables the payment methods, adds the keys and webhook settings in Vercel, tests one booking, then replaces the test credentials with live credentials after approval.",
    )

    doc.add_heading("1. What the client must prepare", level=1)
    add_bullet(doc, "• Business owner’s valid government ID and access to the registered email/phone.")
    add_bullet(doc, "• Business registration and tax documents requested by PayMongo for the chosen account type.")
    add_bullet(doc, "• Bank account details for settlements/payouts, matching the verified business or owner.")
    add_bullet(doc, "• Final website address: https://maddyandcassyrentals-nine.vercel.app (replace this if a custom domain is adopted).")
    add_callout(
        doc,
        "Security action before live launch",
        "Regenerate any secret key that has appeared in a screenshot, chat, email, or public repository. Never place sk_test_…, sk_live_…, or whsk_… values in the website’s frontend code.",
        fill="FFF3DE",
        accent=GOLD,
    )

    doc.add_heading("2. Configure the PayMongo account", level=1)
    add_numbered(doc, "1. Complete identity verification", "Log in to the PayMongo Dashboard and finish the email OTP, liveness check, government ID, and personal information requested on screen.")
    add_numbered(doc, "2. Complete the business profile", "Use the legal name from the registration documents. For the trade name, use “Rental by Maddy & Cassy.” Describe the business as daily rentals of phones, cameras, and related equipment through an online reservation website.")
    add_numbered(doc, "3. Submit business verification", "Upload only the documents PayMongo requests for the selected entity type. Make sure names, addresses, and registration numbers match across all documents.")
    add_numbered(doc, "4. Add settlement details", "Enter the verified bank account in the payout or wallet settings. Double-check the account name and number before saving.")
    add_numbered(doc, "5. Activate payment methods", "Open Payment Channels and request/enable GCash, Cards, and QR Ph. A method will only appear at checkout when PayMongo has made it available to the account.")

    business_heading = doc.add_heading("Suggested business-profile answers", level=2)
    business_heading.paragraph_format.page_break_before = True
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.7)
    for idx, text in enumerate(("Field", "Recommended entry")):
        shade(table.rows[0].cells[idx], ROSE_DARK)
        set_cell_margins(table.rows[0].cells[idx])
        set_run(table.rows[0].cells[idx].paragraphs[0].add_run(text), bold=True, color=WHITE, size=9)
    set_repeat_table_header(table.rows[0])
    entries = [
        ("Trade name", "Rental by Maddy & Cassy"),
        ("Website", "https://maddyandcassyrentals-nine.vercel.app, or the final custom domain"),
        ("Business description", "Online daily rental of smartphones, cameras, and related equipment, with scheduled pickup or delivery."),
        ("Industry", "Choose the closest available rental, retail, electronics, or consumer-services category."),
        ("Contact details", "Use the official business email, mobile number, and complete operating address."),
    ]
    for i, pair in enumerate(entries):
        row = table.add_row()
        prevent_row_split(row)
        for j, text in enumerate(pair):
            shade(row.cells[j], WHITE if i % 2 == 0 else BLUSH_LIGHT)
            set_cell_margins(row.cells[j])
            set_run(row.cells[j].paragraphs[0].add_run(text), size=9, color=INK, bold=(j == 0))
    border = {"val": "single", "sz": 4, "space": 0, "color": LINE}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)

    doc.add_heading("3. Connect test mode to the website", level=1)
    add_numbered(doc, "1. Switch to Test Mode", "Use the Test/Live toggle in the PayMongo Dashboard.")
    add_numbered(doc, "2. Copy test API keys", "Open Developers → API Keys. Copy the test public key and test secret key. Keep the secret key private.")
    add_numbered(doc, "3. Add the values in Vercel", "Open the Vercel project → Settings → Environment Variables. Add these values to Production, Preview, and Development only as needed, then save.")
    add_env_table(doc)
    add_numbered(doc, "4. Create the test webhook", "In PayMongo Test Mode, open Developers → Webhooks → Create Webhook. Use the endpoint below and subscribe to checkout_session.payment.paid.")
    p = doc.add_paragraph(style="Code Block")
    shade_container = OxmlElement("w:shd")
    shade_container.set(qn("w:fill"), "F5F1F0")
    p._p.get_or_add_pPr().append(shade_container)
    set_run(p.add_run("https://maddyandcassyrentals-nine.vercel.app/api/paymongo/webhook"), font="Consolas", size=8.7, color=INK)
    add_numbered(doc, "5. Copy the webhook secret", "Add the generated whsk_… value to PAYMONGO_WEBHOOK_SECRET in Vercel.")
    add_numbered(doc, "6. Redeploy", "Vercel must create a new deployment after environment variables change. Open Deployments, use the latest deployment menu, and select Redeploy—or push a harmless new commit.")

    doc.add_heading("4. Test the complete payment flow", level=1)
    checks = [
        "Open the deployed website, create or sign in to a customer account, and start a new reservation.",
        "Choose 50% or full payment and click Continue to PayMongo.",
        "Select an available test payment method. For e-wallet tests, choose Authorize on PayMongo’s simulation page.",
        "Return to the merchant. The booking must resume instead of returning to a blank or login-only state.",
        "Confirm the booking shows Payment received/verified and that a receipt appears under the customer’s booking.",
        "Open PayMongo → Developers → Webhooks and confirm the event received a successful response.",
    ]
    for item in checks:
        add_bullet(doc, f"☐ {item}")
    add_callout(
        doc,
        "Important QR Ph test warning",
        "PayMongo’s current testing documentation warns that test-mode QR Ph can generate a real QR code. Do not scan and pay it. Use PayMongo’s provided test simulation link instead.",
        fill="FFF3DE",
        accent=GOLD,
    )

    doc.add_heading("5. Move from test to live", level=1)
    add_numbered(doc, "1. Wait for activation", "Confirm business verification is approved and the required payment channels are active in Live Mode.")
    add_numbered(doc, "2. Regenerate live credentials", "Open Developers → API Keys in Live Mode. Rotate any live secret that was previously exposed.")
    add_numbered(doc, "3. Create a separate live webhook", "Create it while the dashboard is in Live Mode, using the same production endpoint and checkout_session.payment.paid event. Test and live webhooks are separate.")
    add_numbered(doc, "4. Replace production values in Vercel", "Use pk_live_…, sk_live_…, and the live webhook’s whsk_…. Do not retain test keys in Production.")
    add_numbered(doc, "5. Redeploy and run a small real payment", "Use a low-value booking, verify the PayMongo payment, webhook delivery, receipt, and booking status, then refund or reconcile it according to the business process.")
    add_callout(doc, "Live launch rule", "A PayMongo checkout page opening is not enough. Launch only after the website receives the paid webhook and records the correct booking and receipt.", fill=GREEN_LIGHT, accent=GREEN)

    doc.add_heading("6. Tutorial videos and visual references", level=1)
    p = doc.add_paragraph()
    set_run(p.add_run("These videos are useful for seeing the flow. PayMongo’s current official documentation, linked below, is the final reference because dashboard labels and API behavior can change."), color=GRAY, italic=True, size=9.5)

    add_video(
        doc,
        1,
        "PayMongo API Tutorial Part 2 — Registering to PayMongo [Tagalog]",
        "https://www.youtube.com/watch?v=YUGp07cKUlg",
        "Best for visually understanding account creation and initial registration.",
        "Published in 2021; the current dashboard may look different.",
    )
    add_video(
        doc,
        2,
        "PayMongo API Tutorial Part 3 — Creating a Payment Source [Tagalog]",
        "https://www.youtube.com/watch?v=rS3QuyONcvA",
        "Shows the customer redirection/authorization idea behind an e-wallet payment.",
        "Conceptual only for this project; Rental by Maddy & Cassy uses PayMongo Hosted Checkout Sessions.",
    )
    add_video(
        doc,
        3,
        "PayMongo API Tutorial Part 4 — Webhook and Payment [Tagalog]",
        "https://www.youtube.com/watch?v=ZEh7UtvI-K0",
        "Useful for understanding why the webhook—not the browser return page—confirms a payment.",
        "Published in 2021; follow the current dashboard webhook steps in this guide.",
    )
    add_video(
        doc,
        4,
        "Online Payment Method Using PHP and PayMongo",
        "https://www.youtube.com/watch?v=g8Pt2KBusvE",
        "A newer visual example of checkout plus real-time webhook handling.",
        "The programming language differs from this Next.js system, so use it only to visualize the payment lifecycle.",
    )

    doc.add_heading("Current official visual guides", level=2)
    links = [
        ("Create and verify a PayMongo account", "https://docs.paymongo.com/docs/get-started-create-your-account"),
        ("Hosted Checkout quick start", "https://docs.paymongo.com/docs/payment-channels-hosted-checkout-quick-start"),
        ("Create a webhook in the dashboard", "https://docs.paymongo.com/docs/developer-tools-dashboard-module-create-a-webhook"),
        ("Test cards, e-wallets, and QR Ph", "https://docs.paymongo.com/docs/payment-acceptance-testing"),
        ("PayMongo Dashboard Mastery Hub", "https://paymongo.help/en/collections/3185124-paymongo-dashboard-mastery-hub"),
    ]
    for title, url in links:
        p = doc.add_paragraph(style="Compact Bullet")
        set_run(p.add_run("• "), color=ROSE_DARK)
        add_hyperlink(p, title, url)

    doc.add_heading("7. Fast troubleshooting", level=1)
    trouble = [
        ("No payment methods appear", "Check Live/Test mode, channel activation, PAYMONGO_PAYMENT_METHODS, and redeploy Vercel."),
        ("Secure checkout cannot be created", "Confirm the server secret key matches the current mode and contains no spaces or quotation marks."),
        ("Payment succeeds but booking stays pending", "Check the webhook mode, endpoint, event selection, secret, and latest Vercel function logs."),
        ("Customer returns to login", "Confirm NEXT_PUBLIC_APP_URL matches the deployed domain and that the user is returning to the same domain used to start the booking."),
        ("Webhook shows failed delivery", "Open the event in PayMongo, inspect the response, fix the endpoint, then resend the event if the dashboard offers that action."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.15)
    table.columns[1].width = Inches(4.35)
    for idx, text in enumerate(("Problem", "What to check")):
        shade(table.rows[0].cells[idx], ROSE_DARK)
        set_cell_margins(table.rows[0].cells[idx])
        set_run(table.rows[0].cells[idx].paragraphs[0].add_run(text), bold=True, color=WHITE, size=9)
    set_repeat_table_header(table.rows[0])
    for i, pair in enumerate(trouble):
        row = table.add_row()
        prevent_row_split(row)
        for j, text in enumerate(pair):
            shade(row.cells[j], WHITE if i % 2 == 0 else BLUSH_LIGHT)
            set_cell_margins(row.cells[j])
            set_run(row.cells[j].paragraphs[0].add_run(text), size=9, color=INK, bold=(j == 0))
    border = {"val": "single", "sz": 4, "space": 0, "color": LINE}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)

    doc.add_heading("Client handoff checklist", level=1)
    for item in (
        "PayMongo identity and business verification approved",
        "Settlement bank account confirmed",
        "GCash, Cards, and QR Ph status reviewed",
        "Test keys and test webhook configured in Vercel",
        "Full test booking completed and recorded",
        "Exposed secret keys rotated",
        "Live keys and live webhook configured",
        "Small live transaction verified end to end",
    ):
        add_bullet(doc, f"☐ {item}")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Keep all secret credentials in PayMongo and Vercel—not in chat, screenshots, or source control."), bold=True, color=ROSE_DARK, size=9.5)

    doc.add_heading("Launch sign-off record", level=2)
    signoff = doc.add_table(rows=4, cols=2)
    signoff.alignment = WD_TABLE_ALIGNMENT.CENTER
    signoff.autofit = False
    signoff.columns[0].width = Inches(3.25)
    signoff.columns[1].width = Inches(3.25)
    labels = [
        ("Client/merchant representative", "Technical representative"),
        ("PayMongo verification approved on", "Live payment channels confirmed"),
        ("Test booking reference", "Live verification transaction reference"),
        ("Client approval/signature", "Date"),
    ]
    for row, pair in zip(signoff.rows, labels):
        prevent_row_split(row)
        for idx, label in enumerate(pair):
            cell = row.cells[idx]
            set_cell_width(cell, 3.25)
            shade(cell, BLUSH_LIGHT)
            set_cell_margins(cell, top=100, start=130, bottom=230, end=130)
            set_run(cell.paragraphs[0].add_run(label), bold=True, color=ROSE_DARK, size=8.8)
    border = {"val": "single", "sz": 5, "space": 0, "color": LINE}
    for row in signoff.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)

    doc.add_heading("After launch: simple monthly checks", level=2)
    for item in (
        "Review failed payments and failed webhook deliveries.",
        "Reconcile PayMongo settlements against recorded booking receipts.",
        "Confirm the live payment channels still appear at checkout.",
        "Remove access for former staff and rotate a key immediately if it may have been exposed.",
    ):
        add_bullet(doc, f"• {item}")

    # Set a useful document title and refresh fields when opened in Word.
    doc.core_properties.title = "PayMongo Client Configuration Guide"
    doc.core_properties.subject = "Rental by Maddy & Cassy payment setup"
    doc.core_properties.author = "Rental by Maddy & Cassy"
    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
